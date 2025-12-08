# servidor.py  (VERSIÓN ACTUALIZADA - Dinámico, admite actividades e imágenes con tokens flexibles)
import os
import tempfile
import shutil
import json
import re
from datetime import datetime
from flask import Flask, request, send_file, jsonify, render_template
from flask_cors import CORS
import pandas as pd
import unicodedata
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__, template_folder="templates")
CORS(app)


# ------------------------
# Utilitarios (tu código original, sin cambios lógicos)
# ------------------------
def normalizar(texto):
    if texto is None:
        return ""
    return (
        unicodedata.normalize("NFKD", str(texto).lower())
        .encode("ascii", "ignore")
        .decode("ascii")
        .replace(" ", "_")
    )


def valOrDash(v):
    return v if (v is not None and str(v).strip() != "") else "-"


def set_cell_style(cell, text, font_size=10, bold=False, align_center=True):
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(font_size)
            run.bold = bold
        paragraph.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
            if align_center
            else WD_PARAGRAPH_ALIGNMENT.LEFT
        )
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_subtitle(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)


def add_note(doc, note="*NO CUENTA CON DICHO ELEMENTO"):
    p = doc.add_paragraph()
    run = p.add_run(note)
    run.italic = True
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.name = "Calibri"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_after = Pt(6)


def create_table(doc, rows, cols, font_size=10, indent=False):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if indent:
        tbl = table._element  # <w:tbl>
        tblPr_elems = tbl.xpath("w:tblPr")
        if tblPr_elems:
            tblPr = tblPr_elems[0]
        else:
            tblPr = OxmlElement("w:tblPr")
            tbl.append(tblPr)
        tblInd = OxmlElement("w:tblInd")
        tblInd.set(qn("w:w"), "300")
        tblInd.set(qn("w:type"), "dxa")
        tblPr.append(tblInd)
    for row in table.rows:
        for cell in row.cells:
            set_cell_style(cell, "-", font_size=font_size)
    return table


def insertar_recuadro_foto(doc, ancho_cm=15, alto_cm=10):
    """
    Inserta un recuadro placeholder (si no hay imagen).
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    try:
        ancho_in = ancho_cm / 2.54
        alto_twips = int(alto_cm * 567)  # 1 cm ≈ 567 twips
    except Exception:
        ancho_in = 15 / 2.54
        alto_twips = int(10 * 567)

    cell = table.cell(0, 0)
    try:
        cell.width = Inches(ancho_in)
    except Exception:
        pass

    # altura exacta de fila
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(alto_twips))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

    # texto central de placeholder
    p = cell.paragraphs[0]
    run = p.add_run("ESPACIO PARA IMAGEN")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # fondo + borde
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "FFFFFF")
    tcPr.append(shd)

    borders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    tcPr.append(borders)


# ------------------------
# Manejo de imágenes (guardado temporal y búsqueda flexible)
# ------------------------
def save_uploaded_files_tmp(request_files):
    """
    Guarda todos los archivos subidos en carpeta temporal y devuelve lista de dicts:
    [{ 'field': <campo form>, 'filename': <original name>, 'path': <tmp path> }, ...]
    """
    saved = []
    tmpdir = tempfile.mkdtemp(prefix="uploaded_imgs_")
    for key in request_files:
        f = request_files.getlist(key) if hasattr(request_files, "getlist") else [request_files[key]]
        for fh in f:
            if fh and getattr(fh, "filename", None):
                safe_name = fh.filename
                tmp_path = os.path.join(tmpdir, safe_name)
                fh.save(tmp_path)
                saved.append({"field": key, "filename": safe_name, "path": tmp_path})
    return saved, tmpdir


def find_images_for_token(images_list, token):
    """
    Busca imágenes cuya filename contenga token (case-insensitive).
    Si token es None o vacío devuelve [].
    """
    if not token:
        return []
    token_l = token.lower()
    out = []
    for item in images_list:
        if token_l in item["filename"].lower():
            out.append(item["path"])
    return out


def find_images_for_any_token(images_list, tokens):
    """
    Dado un listado de tokens (strings), devuelve lista de rutas de imágenes
    cuyo filename contenga cualquiera de los tokens (orden preservado y unicidad).
    """
    if not tokens:
        return []
    found = []
    seen = set()
    for t in tokens:
        if not t:
            continue
        t_l = t.lower()
        for item in images_list:
            if t_l in item["filename"].lower():
                p = item["path"]
                if p not in seen:
                    found.append(p)
                    seen.add(p)
    return found


def insert_images_one_per_line(doc, image_paths, ancho_cm=15, alto_cm=10):
    """
    Inserta una serie de imágenes en el documento, una por línea (una debajo de otra).
    Ajusta cada imagen a ancho_cm x alto_cm (en Inches).
    """
    if not image_paths:
        return
    ancho_in = ancho_cm / 2.54
    alto_in = alto_cm / 2.54
    for p in image_paths:
        try:
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            try:
                table.cell(0, 0).width = Inches(ancho_in)
            except Exception:
                pass

            tr = table.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), str(int(alto_cm * 567)))
            trHeight.set(qn("w:hRule"), "exact")
            trPr.append(trHeight)

            cell = table.cell(0, 0)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            try:
                run.add_picture(p, width=Inches(ancho_in), height=Inches(alto_in))
            except Exception:
                try:
                    doc.add_picture(p, width=Inches(ancho_in), height=Inches(alto_in))
                except Exception:
                    cell.text = "ESPACIO PARA IMAGEN (ERROR AL INSERTAR)"
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "FFFFFF")
            tcPr.append(shd)
            borders = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "12")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "000000")
                borders.append(border)
            tcPr.append(borders)
        except Exception:
            insertar_recuadro_foto(doc, ancho_cm=ancho_cm, alto_cm=alto_cm)


# ------------------------
# Función central: genera docx desde DataFrames + actividades + lista de imágenes temporales
# ------------------------
def generar_docx_desde_dfs(
    df_info, df_tanques, df_accesorios, df_red, df_equipos, df_obs, actividades_list=None, images_list=None
):
    """
    actividades_list: lista de dicts {id, contexto, titulo, tiempo, estado}
    images_list: lista de dicts {'field','filename','path'}
    """
    actividades_list = actividades_list or []
    images_list = images_list or []
    doc = Document()

    # --- Título
    titulo = doc.add_paragraph()
    run = titulo.add_run(
        "INFORME DE MANTENIMIENTO PREVENTIVO Y CUMPLIMIENTO NORMATIVO"
    )
    run.bold = True
    run.underline = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # === 1. Información de cliente
    add_subtitle(doc, "1. INFORMACIÓN DE CLIENTE")
    campos = [
        "Nombre o razón social del cliente",
        "Fecha de inspección",
        "Dirección",
        "RUC o DNI",
        "Número de instalación",
        "Distrito",
        "Departamento",
        "Coordenadas",
        "Nombre del contacto",
        "Número del contacto",
        "Correo de contacto",
    ]
    tabla1 = create_table(doc, len(campos), 2)
    datos_generales = {}
    if df_info is not None and not df_info.empty:
        datos_generales = df_info.iloc[0].to_dict()
    for i, campo in enumerate(campos):
        valor = datos_generales.get(campo, None)
        set_cell_style(tabla1.cell(i, 0), campo, align_center=False)
        set_cell_style(tabla1.cell(i, 1), valOrDash(valor), align_center=False)

    # === 2. Tipo de instalacion
    add_subtitle(doc, "2. TIPO DE INSTALACION")
    tabla2 = create_table(doc, 3, 8)
    tabla2.cell(0, 0).merge(tabla2.cell(0, 1)).text = "DOMESTICO"
    tabla2.cell(0, 2).merge(tabla2.cell(0, 5)).text = "INDUSTRIAL"
    tabla2.cell(0, 6).merge(tabla2.cell(0, 7)).text = "CANALIZADO"
    for cell in tabla2.rows[0].cells:
        set_cell_style(cell, cell.text, bold=True)
    subtipos = [
        "Doméstico",
        "Comercio",
        "Industrial",
        "Agroindustrial",
        "Minera",
        "Avícola",
        "Residencial",
        "Comercial",
    ]
    for i, subtipo in enumerate(subtipos):
        set_cell_style(tabla2.cell(1, i), subtipo)
    for i in range(8):
        set_cell_style(tabla2.cell(2, i), " ")

    # === 3. Tanques inspeccionados
    add_subtitle(doc, "3. TANQUES INSPECCIONADOS")
    headers3 = [
        "Tanque",
        "Capacidad (gal)",
        "N° de serie",
        "Año de fabricación",
        "Tipo de tanque",
        "Fabricante de Tanque",
        "% Actual",
    ]
    num_tanques = len(df_tanques) if df_tanques is not None else 0
    tabla3 = create_table(doc, num_tanques + 1, len(headers3))
    for j, h in enumerate(headers3):
        set_cell_style(tabla3.cell(0, j), h, bold=True)
    for i in range(num_tanques):
        for j, col in enumerate(headers3):
            valor = None
            if df_tanques is not None and col in df_tanques.columns:
                valor = df_tanques.iloc[i][col]
            else:
                if col == "N° de serie" and df_tanques is not None:
                    if "serie" in df_tanques.columns:
                        valor = df_tanques.iloc[i].get("serie", None)
            if j == 0:
                valor = str(i + 1)
            set_cell_style(tabla3.cell(i + 1, j), valOrDash(valor))

    # === 4. Accesorios de los tanques
    add_subtitle(doc, "4. ACCESORIOS DE LOS TANQUES")
    accesorios = [
        "Válvula de llenado",
        "Medidor de porcentaje",
        "Válvula de seguridad",
        "Válvula de drenaje",
        "Multiválvula",
        "Válvula exceso de flujo (Retorno)",
        "Válvula exceso de flujo (Bypass)",
        "Val 3",
    ]
    atributos = ["Marca", "Código", "Serie", "Mes/Año de fabricación"]
    if df_accesorios is None or df_accesorios.empty:
        df_accesorios = pd.DataFrame(columns=["Tanque", "Atributo"] + accesorios)
    unique_tanques = (
        sorted(df_accesorios["Tanque"].unique()) if not df_accesorios.empty else []
    )
    nrows = 1 + len(atributos) * len(unique_tanques) if unique_tanques else 2
    tabla4 = doc.add_table(rows=nrows, cols=2 + len(accesorios))
    tabla4.style = "Table Grid"
    tabla4.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_style(tabla4.cell(0, 0), "N", font_size=7, bold=True)
    set_cell_style(tabla4.cell(0, 1), "Tanques", font_size=7, bold=True)
    for i, acc in enumerate(accesorios):
        set_cell_style(tabla4.cell(0, i + 2), acc, font_size=7, bold=True)
    row_idx = 1
    for tanque in unique_tanques:
        grupo = df_accesorios[df_accesorios["Tanque"] == tanque]
        for attr in atributos:
            set_cell_style(tabla4.cell(row_idx, 1), attr, font_size=7)
            for i, acc in enumerate(accesorios):
                try:
                    val = grupo[grupo["Atributo"] == attr][acc].values
                    if len(val) and pd.notna(val[0]) and str(val[0]).strip() != "":
                        valor = val[0]
                    else:
                        valor = "-"
                except Exception:
                    valor = "-"
                set_cell_style(tabla4.cell(row_idx, i + 2), str(valor), font_size=7)
            row_idx += 1
        # merge number column
        if row_idx - 1 >= row_idx - 4:
            tabla4.cell(row_idx - 4, 0).merge(tabla4.cell(row_idx - 1, 0))
            set_cell_style(tabla4.cell(row_idx - 4, 0), str(tanque), font_size=7)

    # === 5. Accesorios en redes ===
    add_subtitle(doc, "5. ACCESORIOS EN REDES")
    df_red_local = (
        df_red.copy()
        if df_red is not None
        else pd.DataFrame(
            columns=["Tipo", "Marca", "Serie", "Código", "Mes/Año de fabricación"]
        )
    )
    if "Tipo" in df_red_local.columns:
        df_red_local["Tipo"] = df_red_local["Tipo"].astype(str).str.lower().fillna("")
    else:
        df_red_local["Tipo"] = ""
    # Mapa de accesorios
    mapa_accesorios = {
        "llenado_toma_desplazada": "5.1. Válvula de llenado (toma desplazada)",
        "retorno_toma_desplazada": "5.2. Válvula de retorno (toma desplazada)",
        "alivio_hidrostatico": "5.3. Válvula de alivio hidrostático",
        "regulador_primera_etapa": "5.4. Regulador de primera etapa",
        "alivio": "5.5. Válvula de alivio",
        "regulador_2da": "5.6. Regulador de segunda etapa",
        "pull_away": "5.7. Válvula Pull Away",
    }
    grupos = df_red_local.groupby("Tipo") if not df_red_local.empty else {}

    accesorios_red_dict = {}
    for clave, titulo in mapa_accesorios.items():
        add_subtitle(doc, titulo, indent=True)
        lista = (
            grupos.get_group(clave).to_dict(orient="records")
            if (hasattr(grupos, "groups") and clave in grupos.groups)
            else []
        )
        filas = max(2, len(lista) + 1)
        tabla = create_table(doc, filas, 5, indent=True)
        headers = [
            "Válvula",
            "Marca",
            "Serie",
            "Código",
            "Mes/Año de fabricación",
        ]
        for j, h in enumerate(headers):
            set_cell_style(tabla.cell(0, j), h, bold=True)
        if lista:
            for idx, acc in enumerate(lista):
                set_cell_style(tabla.cell(idx + 1, 0), str(idx + 1))
                set_cell_style(tabla.cell(idx + 1, 1), valOrDash(acc.get("Marca")))
                set_cell_style(tabla.cell(idx + 1, 2), valOrDash(acc.get("Serie")))
                set_cell_style(tabla.cell(idx + 1, 3), valOrDash(acc.get("Código")))
                set_cell_style(
                    tabla.cell(idx + 1, 4), valOrDash(acc.get("Mes/Año de fabricación"))
                )
        else:
            for j in range(5):
                set_cell_style(tabla.cell(1, j), "-")
        accesorios_red_dict[clave] = lista

    # Zona medidores 
    zona_medidores_bool = False
    try:
        zona_medidores_bool = (
            df_red_local[df_red_local["Tipo"] == "zona_medidores"]["Código"]
            .astype(str)
            .str.lower()
            .str.contains("true")
            .any()
        )
    except Exception:
        zona_medidores_bool = False
    accesorios_red_dict["zona_medidores"] = zona_medidores_bool

    # === 6. Equipos de la instalación ===
    add_subtitle(doc, "6. EQUIPOS DE LA INSTALACIÓN")
    df_equipos_local = df_equipos.copy() if df_equipos is not None else pd.DataFrame()
    if "Tipo de equipo" in df_equipos_local.columns:
        df_equipos_local["Tipo de equipo"] = (
            df_equipos_local["Tipo de equipo"].astype(str).str.lower().fillna("")
        )
    else:
        df_equipos_local["Tipo de equipo"] = ""

    estructura_equipos = {
        "vaporizador": ["Equipo", "Marca", "Tipo", "Serie", "Año de fabricación", "Capacidad"],
        "quemador": ["Equipo", "Marca", "Modelo", "Tipo", "Serie", "Año de fabricación", "Capacidad (kW)"],
        "decantador": ["Equipo", "Fabricante", "Modelo", "Tipo", "Serie", "Año de fabricación", "Capacidad (gal)"],
        "dispensador_de_gas": ["Equipo", "Marca", "Modelo", "Serie"],
        "bomba": ["Equipo", "Marca", "Modelo", "Serie"],
        "tablero": ["Equipo", "TAG"],
        "estabilizador": ["Equipo", "Marca", "Modelo", "Serie"],
        "detector": ["Equipo", "Marca", "Modelo", "Serie"],
        "extintor": ["Equipo", "Marca", "Serie", "Año de fabricación", "Próxima PH", "Fecha de próxima recarga"],
    }

    grupos_e = df_equipos_local.groupby("Tipo de equipo") if not df_equipos_local.empty else {}

    for idx, (tipo_equipo, columnas) in enumerate(estructura_equipos.items(), start=1):
        nombre_limpio = tipo_equipo.replace("_", " ").capitalize()
        subtitulo = f"6.{idx}. {nombre_limpio}"
        add_subtitle(doc, subtitulo, indent=True)
        datos = (
            grupos_e.get_group(tipo_equipo)
            if (hasattr(grupos_e, "groups") and tipo_equipo in grupos_e.groups)
            else pd.DataFrame(columns=columnas)
        )
        filas = max(2, len(datos) + 1)
        tabla = create_table(doc, filas, len(columnas), indent=True)
        for j, col in enumerate(columnas):
            set_cell_style(tabla.cell(0, j), col, bold=True)
        if not datos.empty:
            for i, (_, fila) in enumerate(datos.iterrows()):
                for j, col in enumerate(columnas):
                    valor = fila.get(col, None)
                    if j == 0:
                        set_cell_style(tabla.cell(i + 1, j), str(i + 1))
                    else:
                        set_cell_style(tabla.cell(i + 1, j), valOrDash(valor))
        else:
            for j in range(len(columnas)):
                set_cell_style(tabla.cell(1, j), "-")

    equipos_instalacion = {
        k: grupos_e.get_group(k).to_dict(orient="records")
        if (hasattr(grupos_e, "groups") and k in grupos_e.groups)
        else []
        for k in estructura_equipos.keys()
    }

    # === 7. Observaciones generales ===
    add_subtitle(doc, "7. OBSERVACIONES GENERALES")
    df_obs_local = df_obs.copy() if df_obs is not None else pd.DataFrame(columns=["Subpunto", "Observación"])
    subtitulos_7 = {
        "7.1": "7.1. Observaciones al cliente",
        "7.2": "7.2. Observaciones en red de llenado y retorno",
        "7.3": "7.3. Observaciones en zona de tanque",
        "7.4": "7.4. Observaciones en red de consumo",
    }
    for clave, titulo in subtitulos_7.items():
        add_subtitle(doc, titulo, indent=True)
        texto = df_obs_local[df_obs_local["Subpunto"] == clave]["Observación"].values
        if len(texto) and str(texto[0]).strip() != "":
            doc.add_paragraph(str(texto[0]).strip())
        else:
            doc.add_paragraph("-")

    add_subtitle(doc, "7.5. Observaciones en equipos varios (Vaporizador, Quemador, Decantador, etc)", indent=True)
    equipos_obs = [
        "Vaporizador",
        "Quemador",
        "Decantador",
        "Dispensador de gas",
        "Bomba de abastecimiento",
        "Tablero eléctrico",
        "Estabilizador",
        "Detector de gases",
        "Extintor",
    ]
    tabla_obs = create_table(doc, len(equipos_obs) + 1, 2, indent=True)
    set_cell_style(tabla_obs.cell(0, 0), "Equipo", bold=True)
    set_cell_style(tabla_obs.cell(0, 1), "Observación", bold=True)
    texto_75 = df_obs_local[df_obs_local["Subpunto"] == "7.5"]["Observación"].values
    observaciones_75 = []
    if len(texto_75) and str(texto_75[0]).strip():
        observaciones_75 = [x.strip() for x in str(texto_75[0]).split(".") if x.strip()]
    for i, equipo in enumerate(equipos_obs):
        set_cell_style(tabla_obs.cell(i + 1, 0), equipo)
        set_cell_style(tabla_obs.cell(i + 1, 1), observaciones_75[i] if i < len(observaciones_75) else "-")

    # === 8. Evidencia general ===
    add_subtitle(doc, "8. EVIDENCIA FOTOGRÁFICA (del establecimiento)")
    # tokens que el frontend puede usar: 'sub_8_establecimiento' o '8' o '8_establecimiento'
    imgs_8 = find_images_for_any_token(images_list, ["sub_8_establecimiento", "8_establecimiento", "8"])
    if imgs_8:
        insert_images_one_per_line(doc, imgs_8, ancho_cm=15, alto_cm=10)
    else:
        insertar_recuadro_foto(doc)

    # === 9. Evidencia fotográfica de elementos de la instalación ===
    add_subtitle(doc, "9. Evidencia fotográfica de elementos de la instalación")

    # Construyo un bloque flexible: intentaré encontrar imágenes por tokens numéricos (9_1, 9_2...) y por tokens descriptivos que usa el frontend
    tanques_for_block = (
        df_tanques.to_dict(orient="records")
        if df_tanques is not None and not df_tanques.empty
        else []
    )
    equipos_for_block = (
        df_equipos_local.to_dict(orient="records")
        if df_equipos_local is not None and not df_equipos_local.empty
        else []
    )
    accesorios_red_for_block = accesorios_red_dict

    bloque_9 = []
    contador = 1

    # 9.1 Panorámica general (id numeric '9_{contador}' and also '9_panoramica_general' and '9_panoramica')
    bloque_9.append((f"9.{contador}. FOTO PANORÁMICA DE LA ZONA", True, 1, [f"9_{contador}", "9_panoramica_general", "9_panoramica", "9_panoramica_zona"])); contador += 1

    # Placas por tanque -> tokens: numeric '9_{n}' and descriptive 'tanque_{i}_placa'
    for i, t in enumerate(tanques_for_block):
        serie = valOrDash(t.get("N° de serie") or t.get("serie"))
        tokens = [f"9_{contador}", f"tanque_{i+1}_placa", f"tanque_{i+1}__placa"]
        bloque_9.append((f"9.{contador}. PLACA DE TANQUE {i+1} DE SERIE: {serie}", True, 1, tokens)); contador += 1

    # Panorámica de alrededores por tanque (4 recuadros). tokens: tanque_{i}_panoramica
    for i, t in enumerate(tanques_for_block):
        serie = valOrDash(t.get("N° de serie") or t.get("serie"))
        tokens = [f"9_{contador}", f"tanque_{i+1}_panoramica", f"tanque_{i+1}__panoramica"]
        bloque_9.append((f"9.{contador}. FOTO PANORÁMICA DE ALREDEDORES DE TANQUE {i+1} DE SERIE: {serie}", True, 4, tokens)); contador += 1

    # Bloque iterativo por tanque: varias fotos específicas
    for i, t in enumerate(tanques_for_block):
        serie = valOrDash(t.get("N° de serie") or t.get("serie"))
        items_per_tank = [
            ("FOTO DE BASES DE CONCRETO", f"tanque_{i+1}_bases"),
            ("FOTO DE MANÓMETROS 0-60 PSI", f"tanque_{i+1}_manometro_0_60"),
            ("FOTO DE MANÓMETROS 0-300 PSI", f"tanque_{i+1}_manometro_0_300"),
            ("FOTO DE CONEXIÓN DE CHICOTE A LA MULTIVÁLVULA", f"tanque_{i+1}_chicote"),
            ("STICKERS DEL TANQUE Y PINTADO", f"tanque_{i+1}_stickers"),
            ("FOTO DE LOS 04 ANCLAJES, PERNOS, TORNILLOS", f"tanque_{i+1}_anclajes"),
            ("FOTO DE VÁLVULA DE LLENADO", f"tanque_{i+1}_valvula_llenado"),
            ("FOTO DE VÁLVULA DE SEGURIDAD", f"tanque_{i+1}_valvula_seguridad"),
            ("FOTO DE VÁLVULA DE DRENAJE", f"tanque_{i+1}_valvula_drenaje"),
            ("FOTO DE MULTIVÁLVULA", f"tanque_{i+1}_multivalvula"),
            ("FOTO DE MEDIDOR DE PORCENTAJE", f"tanque_{i+1}_medidor_porcentaje"),
        ]
        for name, token_base in items_per_tank:
            tokens = [f"9_{contador}", token_base, f"tanque_{i+1}__{token_base}"]
            bloque_9.append((f"9.{contador}. {name} DE TANQUE {i+1} DE SERIE: {serie}", True, 1, tokens))
            contador += 1

    # Equipos específicos: intentamos tokens tipo 'equipo_{i}_placa' o '9_{n}'
    tipos_to_iter = ["estabilizador", "quemador", "vaporizador", "tablero", "bomba", "dispensador_de_gas", "decantador", "detector"]
    for tipo in tipos_to_iter:
        lista_eq = equipos_instalacion.get(tipo, [])
        if lista_eq:
            for eq in lista_eq:
                serie = valOrDash(eq.get("Serie"))
                tokens = [f"9_{contador}", f"{tipo}_placa", f"{tipo}_foto", f"equipo_{serie}_placa"]
                bloque_9.append((f"9.{contador}. FOTO DE PLACA DE {tipo.upper()} DE SERIE: {serie}", True, 1, tokens)); contador += 1
                bloque_9.append((f"9.{contador}. FOTO DE {tipo.upper()}", True, 1, [f"9_{contador}", f"{tipo}_foto"])); contador += 1
        else:
            bloque_9.append((f"9.{contador}. FOTO DE PLACA DE {tipo.upper()} DE SERIE: -", False, 1, [f"9_{contador}"])); contador += 1
            bloque_9.append((f"9.{contador}. FOTO DE {tipo.upper()}", False, 1, [f"9_{contador}"])); contador += 1

    # Toma desplazada (llenado_toma_desplazada)
    tiene_toma = bool(accesorios_red_for_block.get("llenado_toma_desplazada"))
    bloque_9.append((f"9.{contador}. FOTO DEL PUNTO DE TRANSFERENCIA DESPLAZADO", tiene_toma, 1, [f"9_{contador}", "toma_desplazada", "llenado_toma_desplazada"])); contador += 1
    bloque_9.append((f"9.{contador}. FOTO DE LA CAJA DE LA TOMA DESPLAZADA", tiene_toma, 1, [f"9_{contador}", "toma_desplazada_caja"])); contador += 1
    bloque_9.append((f"9.{contador}. FOTO DEL RECORRIDO DESDE TOMA DESPLAZADA HASTA TANQUE", tiene_toma, 1, [f"9_{contador}", "toma_desplazada_recorrido"])); contador += 1

    # Accesorios individuales (mapa)
    mapa = {
        "llenado_toma_desplazada": "VÁLVULA DE LLENADO TOMA DESPLAZADA",
        "retorno_toma_desplazada": "VÁLVULA DE RETORNO TOMA DESPLAZADA",
        "alivio": "VÁLVULA DE ALIVIO",
        "regulador_2da": "REGULADOR DE SEGUNDA ETAPA",
        "pull_away": "VÁLVULA PULL AWAY",
        "alivio_hidrostatico": "VÁLVULA DE ALIVIO HIDROSTÁTICO",
        "regulador_primera_etapa": "REGULADOR DE PRIMERA ETAPA",
    }
    for clave, nombre in mapa.items():
        lista = accesorios_red_for_block.get(clave, []) if accesorios_red_for_block else []
        cantidad = max(1, len(lista))
        for idx in range(cantidad):
            if idx < len(lista):
                codigo = valOrDash(lista[idx].get("Código"))
                existe = True
            else:
                codigo = "-"
                existe = False
            tokens = [f"9_{contador}", f"{clave}_{idx+1}", clave]
            bloque_9.append((f"9.{contador}. FOTO DE {nombre} {idx+1} DE CÓDIGO: {codigo}", existe, 1, tokens))
            contador += 1

    # Zona de medidores
    bloque_9.append((f"9.{contador}. FOTO DE ZONA MEDIDORES", bool(accesorios_red_for_block.get("zona_medidores")), 1, [f"9_{contador}", "zona_medidores", "zona_medidor"])); contador += 1

    # Función auxiliar para insertar subtítulos + imgs buscando por tokens alternativos
    def add_foto_con_subtitulo_with_tokens(doc, texto, candidate_tokens, incluir_imagen=True, num_recuadros=1):
        add_subtitle(doc, texto, indent=True)
        if incluir_imagen:
            imgs = find_images_for_any_token(images_list, candidate_tokens if isinstance(candidate_tokens, (list,tuple)) else [candidate_tokens])
            if imgs:
                imgs_to_use = imgs[:num_recuadros] if num_recuadros <= len(imgs) else imgs
                insert_images_one_per_line(doc, imgs_to_use, ancho_cm=15, alto_cm=10)
                if len(imgs_to_use) < num_recuadros:
                    for _ in range(num_recuadros - len(imgs_to_use)):
                        insertar_recuadro_foto(doc)
            else:
                for _ in range(num_recuadros):
                    insertar_recuadro_foto(doc)
        else:
            add_note(doc, "*NO CUENTA CON DICHO ELEMENTO")

    # Recorrer bloque_9 y añadir al doc
    for item in bloque_9:
        texto, incluir, num_recuadros, tokens = item
        if incluir:
            if num_recuadros > 1:
                add_foto_con_subtitulo_with_tokens(doc, texto, tokens, incluir_imagen=True, num_recuadros=num_recuadros)
            else:
                add_foto_con_subtitulo_with_tokens(doc, texto, tokens, incluir_imagen=True, num_recuadros=1)
        else:
            add_foto_con_subtitulo_with_tokens(doc, texto, tokens, incluir_imagen=False, num_recuadros=1)

    # === 10. EVIDENCIA FOTOGRÁFICA (MANTENIMIENTO REALIZADO) ===
    add_subtitle(doc, "10. EVIDENCIA FOTOGRÁFICA (MANTENIMIENTO REALIZADO)")
    add_note(doc, "NOTA 1: SE DEBERÁ MENCIONAR LOS TRABAJOS EJECUTADOS POR TANQUE (INCLUIR LAS INSPECCIONES QUE SE REALICEN)")
    add_note(doc, "NOTA 2: LAS IMÁGENES DEBEN TENER UN TAMAÑO DE 15CM (LARGO) X 10CM (ALTO) MÁXIMO Y SE DEBERÁ VISUALIZAR CLARAMENTE LOS DATOS RELEVANTES (OBSERVACIONES, DESCRIPCIONES DE ESTADO DE ELEMENTOS, TRABAJO REALIZADO, ETC) DE LOS ELEMENTOS EN LOS TRABAJOS REALIZADOS (TANQUES, ACCESORIOS, REDES)")

    # Agrupar actividades por contexto dinámicamente
    # contexts: tanque_{i}, red_{i}, equipo_{i}, general
    actividades_por_contexto = {}
    for act in actividades_list:
        ctx = act.get("contexto") or "general"
        actividades_por_contexto.setdefault(ctx, []).append(act)

    # Numeración 10.x: seguiremos este orden:
    # - por cada tanque existente, si tiene actividades → 10.1, 10.2...
    # - luego TRABAJOS REALIZADOS EN REDES DE LLENADO Y RETORNO (si hay actividades de tipo red y su red Tipo indica llenado/retorno)
    # - luego TRABAJOS REALIZADOS EN REDES DE CONSUMO (resto de red activities)
    # - luego ACTIVIDADES GENERALES (contexto general)
    sec_idx = 1

    # 1) Tanques
    for t_idx in range(len(tanques_for_block)):
        ctx_key = f"tanque_{t_idx+1}"
        acts = actividades_por_contexto.get(ctx_key, [])
        if not acts:
            continue
        # Subtítulo para el tanque
        add_subtitle(doc, f"10.{sec_idx}. TRABAJOS REALIZADOS EN EL TANQUE {t_idx+1} DE SERIE: {valOrDash(tanques_for_block[t_idx].get('N° de serie') or tanques_for_block[t_idx].get('serie'))}", indent=True)
        # para cada actividad en este tanque, insertar título y antes/después buscando por el id del activity
        for a in acts:
            # activity title
            doc.add_paragraph(f"- {a.get('titulo','Actividad')}. Tiempo: {a.get('tiempo','')}. Estado: {a.get('estado','')}")
            # buscar imágenes con token actividad.id + "_before" / "_after" o solo actividad.id
            aid = a.get("id")
            before_tokens = [f"{aid}_before", f"{aid}__before", f"{aid}_antes", aid]
            after_tokens = [f"{aid}_after", f"{aid}__after", f"{aid}_despues", aid]
            imgs_b = find_images_for_any_token(images_list, before_tokens)
            imgs_a = find_images_for_any_token(images_list, after_tokens)
            # insertar antes
            if imgs_b:
                insert_images_one_per_line(doc, imgs_b, ancho_cm=15, alto_cm=10)
            else:
                insertar_recuadro_foto(doc)
            # insertar despues
            if imgs_a:
                insert_images_one_per_line(doc, imgs_a, ancho_cm=15, alto_cm=10)
            else:
                insertar_recuadro_foto(doc)
        sec_idx += 1

    # 2) Redes: separarlas en llenado/retorno vs consumo con base en tipo del accesorio de red
    # recolectar todas las actividades que tengan contexto 'red_{i}'
    red_acts = []
    for k, acts in actividades_por_contexto.items():
        if str(k).startswith("red_"):
            red_acts.extend(acts)

    # clasificar cada actividad segun el tipo de la red asociada (si existe en df_red_local)
    red_llenado = []
    red_consumo = []
    for act in red_acts:
        ctx = act.get("contexto")
        # Extraer índice
        m = re.match(r"red_(\d+)", str(ctx or ""))
        tipo_red_val = ""
        if m:
            idx = int(m.group(1)) - 1
            if 0 <= idx < len(df_red_local):
                tipo_red_val = str(df_red_local.iloc[idx].get("Tipo","") or "").lower()
        # heurística: si contiene 'llenado' o 'retorno' o 'toma' => llenado/retorno
        if any(s in tipo_red_val for s in ["llenado", "retorno", "toma", "llenado_toma"]):
            red_llenado.append(act)
        else:
            red_consumo.append(act)

    # Si existen actividades de red_llenado
    if red_llenado:
        add_subtitle(doc, f"10.{sec_idx}. TRABAJOS REALIZADOS EN REDES DE LLENADO Y RETORNO", indent=True)
        for a in red_llenado:
            doc.add_paragraph(f"- {a.get('titulo','Actividad')}. Tiempo: {a.get('tiempo','')}. Estado: {a.get('estado','')}")
            aid = a.get("id")
            imgs_b = find_images_for_any_token(images_list, [f"{aid}_before", f"{aid}__before", aid])
            imgs_a = find_images_for_any_token(images_list, [f"{aid}_after", f"{aid}__after", aid])
            if imgs_b:
                insert_images_one_per_line(doc, imgs_b, ancho_cm=15, alto_cm=10)
            else:
                insertar_recuadro_foto(doc)
            if imgs_a:
                insert_images_one_per_line(doc, imgs_a, ancho_cm=15, alto_cm=10)
            else:
                insertar_recuadro_foto(doc)
        sec_idx += 1

    # Si existen actividades de red_consumo
    if red_consumo:
        add_subtitle(doc, f"10.{sec_idx}. TRABAJOS REALIZADOS EN REDES DE CONSUMO", indent=True)
        for a in red_consumo:
            doc.add_paragraph(f"- {a.get('titulo','Actividad')}. Tiempo: {a.get('tiempo','')}. Estado: {a.get('estado','')}")
            aid = a.get("id")
            imgs_b = find_images_for_any_token(images_list, [f"{aid}_before", f"{aid}__before", aid])
            imgs_a = find_images_for_any_token(images_list, [f"{aid}_after", f"{aid}__after", aid])
            if imgs_b:
                insert_images_one_per_line(doc, imgs_b, ancho_cm=15, alto_cm=10)
            else:
                insertar_recuadro_foto(doc)
            if imgs_a:
                insert_images_one_per_line(doc, imgs_a, ancho_cm=15, alto_cm=10)
            else:
                insertar_recuadro_foto(doc)
        sec_idx += 1

    # 3) Actividades generales (contexto 'general')
    general_acts = actividades_por_contexto.get("general", [])
    if general_acts:
        add_subtitle(doc, f"10.{sec_idx}. ACTIVIDADES GENERALES", indent=True)
        for a in general_acts:
            doc.add_paragraph(f"- {a.get('titulo','Actividad')}. Tiempo: {a.get('tiempo','')}. Estado: {a.get('estado','')}")
            aid = a.get("id")
            imgs_b = find_images_for_any_token(images_list, [f"{aid}_before", f"{aid}__before", aid])
            imgs_a = find_images_for_any_token(images_list, [f"{aid}_after", f"{aid}__after", aid])
            if imgs_b:
                insert_images_one_per_line(doc, imgs_b, ancho_cm=15, alto_cm=10)
            else:
                insertar_recuadro_foto(doc)
            if imgs_a:
                insert_images_one_per_line(doc, imgs_a, ancho_cm=15, alto_cm=10)
            else:
                insertar_recuadro_foto(doc)
        sec_idx += 1

    # === 11,12,13 ===
    add_subtitle(doc, "11. EVIDENCIA FOTOGRÁFICA DE LA INSTALACIÓN")
    imgs_11 = find_images_for_any_token(images_list, ["11", "11_evidencia", "11_evidencia_instalacion"])
    if imgs_11:
        insert_images_one_per_line(doc, imgs_11, ancho_cm=15, alto_cm=10)
    else:
        insertar_recuadro_foto(doc)

    add_subtitle(doc, "12. Conclusiones")
    doc.add_paragraph("-")
    add_subtitle(doc, "13. Recomendaciones")
    doc.add_paragraph("-")

    # Guardar docx en archivo temporal
    fd, path = tempfile.mkstemp(prefix="Informe_Mantenimiento_", suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path


# ------------------------
# UTIL: construir DataFrames desde JSON (idéntico a tu original)
# ------------------------
def build_dfs_from_json(payload):
    """
    Construye DataFrames desde el payload JSON esperado.
    Retorna: df_info, df_tanques, df_accesorios, df_red, df_equipos, df_obs
    """
    general = payload.get("general", {}) or {}
    tanques = payload.get("tanques", []) or []
    accesorios_tanque = payload.get("accesoriosTanque", {}) or {}
    accesorios_red = payload.get("accesoriosRed", []) or []
    equipos = payload.get("equipos", []) or []
    observaciones = payload.get("observaciones", {}) or {}

    # df_info
    campos = [
        "Nombre o razón social del cliente",
        "Fecha de inspección",
        "Dirección",
        "RUC o DNI",
        "Número de instalación",
        "Distrito",
        "Departamento",
        "Coordenadas",
        "Nombre del contacto",
        "Número del contacto",
        "Correo de contacto",
    ]
    row_info = {c: general.get(c, "") for c in campos}
    df_info = pd.DataFrame([row_info])

    # df_tanques
    tanques_rows = []
    for t in tanques:
        row = {
            "N° de serie": t.get("serie") or t.get("N° de serie") or "",
            "Capacidad (gal)": t.get("capacidad") or "",
            "Año de fabricación": t.get("anio") or t.get("Año de fabricación") or "",
            "Tipo de tanque": t.get("tipo") or "",
            "Fabricante de Tanque": t.get("fabricante") or "",
            "% Actual": t.get("porcentaje") or "",
        }
        tanques_rows.append(row)
    df_tanques = pd.DataFrame(tanques_rows)

    # df_accesorios: construir filas (Tanque, Atributo, <accesorio columns>)
    accesorios_cols = [
        "Válvula de llenado",
        "Medidor de porcentaje",
        "Válvula de seguridad",
        "Válvula de drenaje",
        "Multiválvula",
        "Válvula exceso de flujo (Retorno)",
        "Válvula exceso de flujo (Bypass)",
        "Val 3",
    ]
    atributos = ["Marca", "Código", "Serie", "Mes/Año de fabricación"]
    rows = []
    # accesorios_tanque esperado: { "1": { "Válvula de llenado": {"Marca":..., "Código":...}, ... }, ... }
    for tank_key, accs in accesorios_tanque.items():
        try:
            tk = int(tank_key)
        except Exception:
            tk = tank_key
        for attr in atributos:
            row = {"Tanque": tk, "Atributo": attr}
            for acc_name in accesorios_cols:
                acc_entry = accs.get(acc_name, {}) if isinstance(accs, dict) else {}
                row_val = acc_entry.get(attr, "") if isinstance(acc_entry, dict) else ""
                row[acc_name] = row_val
            rows.append(row)
    df_accesorios = pd.DataFrame(rows)

    # df_red: list of dicts
    df_red = pd.DataFrame(
        [
            {
                "Tipo": r.get("Tipo", ""),
                "Marca": r.get("Marca", ""),
                "Serie": r.get("Serie", ""),
                "Código": r.get("Código", ""),
                "Mes/Año de fabricación": r.get("Mes/Año de fabricación", ""),
            }
            for r in accesorios_red
        ]
    )

    # df_equipos
    df_equipos = pd.DataFrame(equipos)

    # df_obs
    obs_rows = []
    for sp in ["7.1", "7.2", "7.3", "7.4", "7.5"]:
        obs_rows.append({"Subpunto": sp, "Observación": observaciones.get(sp, "")})
    df_obs = pd.DataFrame(obs_rows)

    return df_info, df_tanques, df_accesorios, df_red, df_equipos, df_obs


# ------------------------
# Endpoint /generar actualizado para multipart/form-data (JSON + imágenes)
# ------------------------
@app.route("/generar", methods=["POST"])
def generar_informe():
    tmp_images_dir = None
    saved_images = []
    try:
        # Si el cliente envía multipart/form-data:
        if request.content_type and "multipart/form-data" in request.content_type:
            # Se espera que haya un campo 'json' con el payload
            data_raw = request.form.get("json") or request.form.get("payload") or None
            if not data_raw:
                try:
                    data_raw = request.get_data(as_text=True)
                except Exception:
                    data_raw = None
            try:
                payload = json.loads(data_raw) if data_raw else {}
            except Exception:
                payload = request.get_json(silent=True) or {}
            # guardar archivos
            saved_images, tmp_images_dir = save_uploaded_files_tmp(request.files)
        else:
            # si no es multipart: cuerpo JSON
            payload = request.get_json() or {}
            saved_images = []

        if not payload:
            return jsonify({"error": "No JSON recibido o body vacío"}), 400

        # Validaciones mínimas
        general = payload.get("general", {}) or {}
        tanques = payload.get("tanques", []) or []

        required_general = [
            "Nombre o razón social del cliente",
            "Fecha de inspección",
            "Dirección",
            "RUC o DNI",
            "Número de instalación",
            "Distrito",
            "Departamento",
            "Coordenadas",
            "Nombre del contacto",
            "Número del contacto",
            "Correo de contacto",
        ]
        missing = [k for k in required_general if not (general.get(k) or "").strip()]
        if missing:
            return (
                jsonify({"error": "Faltan campos obligatorios en 'general'", "missing": missing}),
                400,
            )

        if len(tanques) == 0:
            return jsonify({"error": "Se requiere al menos un tanque en 'tanques'"}), 400

        # Validación accesoriosTanque
        accesorios_tanque = payload.get("accesoriosTanque", {}) or {}
        for tk, accs in accesorios_tanque.items():
            for acc_name, fields in (accs or {}).items():
                if any((fields.get(f) or "").strip() for f in ["Marca", "Código", "Serie", "Mes/Año de fabricación"]):
                    missingf = [f for f in ["Marca", "Código", "Serie", "Mes/Año de fabricación"] if not (fields.get(f) or "").strip()]
                    if missingf:
                        return jsonify({"error": f"En accesoriosTanque.{tk}.{acc_name} faltan campos: {missingf}"}), 400

        # Validación accesoriosRed
        accesorios_red = payload.get("accesoriosRed", []) or []
        for i, r in enumerate(accesorios_red):
            if any((r.get(k) or "").strip() for k in ["Marca", "Serie", "Código", "Mes/Año de fabricación"]):
                if not (r.get("Tipo") or "").strip():
                    return jsonify({"error": f"AccesoriosRed[{i}] tiene campos pero falta 'Tipo'"}), 400

        # Validación equipos
        equipos = payload.get("equipos", []) or []
        estructura_equipos = {
            "vaporizador": ["Equipo", "Marca", "Tipo", "Serie", "Año de fabricación", "Capacidad"],
            "quemador": ["Equipo", "Marca", "Modelo", "Tipo", "Serie", "Año de fabricación", "Capacidad (kW)"],
            "decantador": ["Equipo", "Fabricante", "Modelo", "Tipo", "Serie", "Año de fabricación", "Capacidad (gal)"],
            "dispensador_de_gas": ["Equipo", "Marca", "Modelo", "Serie"],
            "bomba": ["Equipo", "Marca", "Modelo", "Serie"],
            "tablero": ["Equipo", "TAG"],
            "estabilizador": ["Equipo", "Marca", "Modelo", "Serie"],
            "detector": ["Equipo", "Marca", "Modelo", "Serie"],
            "extintor": ["Equipo", "Marca", "Serie", "Año de fabricación", "Próxima PH", "Fecha de próxima recarga"],
        }
        for i, eq in enumerate(equipos):
            tipo = (eq.get("Tipo de equipo") or eq.get("tipo") or "").strip().lower()
            if tipo and tipo in estructura_equipos:
                required_cols = estructura_equipos[tipo]
                missing_eq = [c for c in required_cols if not (eq.get(c) or "").strip()]
                if missing_eq:
                    return jsonify({"error": f"Equipo[{i}] de tipo '{tipo}' faltan campos: {missing_eq}"}), 400

        # Construir DataFrames
        df_info, df_tanques, df_accesorios, df_red, df_equipos, df_obs = build_dfs_from_json(payload)

        # Actividades (pueden venir en payload.actividades)
        actividades = payload.get("actividades", []) or []

        # Generar docx (pasando lista de imágenes guardadas y actividades)
        ruta = generar_docx_desde_dfs(df_info, df_tanques, df_accesorios, df_red, df_equipos, df_obs, actividades_list=actividades, images_list=saved_images)

        # Enviar archivo
        try:
            response = send_file(ruta, as_attachment=True, download_name=os.path.basename(ruta))
            return response
        except TypeError:
            return send_file(ruta, as_attachment=True)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": "Error interno del servidor", "detail": str(e)}), 500
    finally:
        # limpieza: remover archivos temporales si se crearon
        try:
            if tmp_images_dir and os.path.isdir(tmp_images_dir):
                shutil.rmtree(tmp_images_dir, ignore_errors=True)
        except Exception:
            pass


# Página simple para probar manualmente
@app.route("/")
def index():
    try:
        return render_template("pagina.html")
    except Exception:
        return "<h3>Servidor Flask funcionando. Envia POST JSON a /generar</h3>"


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
